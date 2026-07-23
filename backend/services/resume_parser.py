import io
import magic
from typing import Tuple, Optional

import pdfplumber
from docx import Document
import PyPDF2

try:
    from backend.utils.file_utils import (
        FileParsingError,
        TextExtractionError,
        FileUploadError,
        log_error,
        log_warning,
        log_info,
        with_fallback,
        ATSBaseError as _ATSBaseError,
    )
    from backend.core.config import (
        MAX_FILE_SIZE_BYTES,
        MAX_FILE_SIZE_MB,
        SUPPORTED_MIME_TYPES,
    )
except ImportError:
    from utils.file_utils import (
        FileParsingError,
        TextExtractionError,
        FileUploadError,
        log_error,
        log_warning,
        log_info,
        with_fallback,
        ATSBaseError as _ATSBaseError,
    )
    from core.config import (
        MAX_FILE_SIZE_BYTES,
        MAX_FILE_SIZE_MB,
        SUPPORTED_MIME_TYPES,
    )


class FileValidationError(_ATSBaseError):
    """Raised when a file fails size / MIME-type / extension validation."""
    pass


# ── File validation ────────────────────────────────────────────────────────────
def validate_file(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
    file_size_bytes = len(file_data)

    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f'File size ({size_mb:.2f} MB) exceeds the {MAX_FILE_SIZE_MB} MB limit. '
            'Please upload a smaller file or compress your resume.'
        ), None

    if file_size_bytes == 0:
        return False, 'The uploaded file is empty. Please check the file and try again.', None

    try:
        mime_type = magic.from_buffer(file_data, mime=True)
    except Exception as e:
        import mimetypes
        guess, _ = mimetypes.guess_type(filename)
        if guess:
            mime_type = guess
        elif filename.lower().endswith('.pdf'):
            mime_type = 'application/pdf'
        elif filename.lower().endswith('.docx'):
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.doc'):
            mime_type = 'application/msword'
        else:
            return False, f'Could not determine the file type: {e}', None

    if mime_type not in SUPPORTED_MIME_TYPES:
        # Try extension fallback in case MIME detection was imprecise
        ext = '.' + filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        ext_map = {
            '.pdf':  'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc':  'application/msword',
        }
        if ext in ext_map:
            mime_type = ext_map[ext]
        else:
            return False, (
                f'Unsupported file type ({mime_type}). '
                'Please upload a PDF or DOCX resume.'
            ), None

    return True, '', SUPPORTED_MIME_TYPES[mime_type]


# ── PDF extraction ─────────────────────────────────────────────────────────────
def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page:
                continue
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    if annot.get('/Subtype') != '/Link':
                        continue
                    uri = annot.get('/A', {}).get('/URI', '')
                    if isinstance(uri, bytes):
                        uri = uri.decode('utf-8', errors='ignore')
                    uri = uri.strip() if isinstance(uri, str) else ''
                    if uri.startswith('http'):
                        urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(urls)


def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text = ''
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.',
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks
    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text = ''
    reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.',
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks
    return text.strip()


def extract_text_from_pdf(file_data: bytes) -> str:
    try:
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            log_fallback=True,
        )
        if used_fallback:
            log_info('PDF extraction succeeded using PyPDF2 fallback', context='resume_parser')
        return result
    except Exception as e:
        log_error(e, context='extract_text_from_pdf')
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e


# ── DOCX extraction ────────────────────────────────────────────────────────────
def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc        = Document(io.BytesIO(file_data))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = '\n'.join(text_parts)

        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )

        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            pass

        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        return text.strip()

    except FileParsingError:
        raise

    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Try re-saving as DOCX or converting to PDF.'
        ) from e


def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert to .docx or .pdf and try again.'
    )


# ── Dispatch ───────────────────────────────────────────────────────────────────
def extract_text(file_data: bytes, file_type: str) -> str:
    if file_type == 'pdf':
        return extract_text_from_pdf(file_data)
    elif file_type == 'docx':
        return extract_text_from_docx(file_data)
    elif file_type == 'doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f'Unsupported file type: {file_type}. Supported: pdf, docx.'
        )


# ── Entry point ────────────────────────────────────────────────────────────────
def parse_resume_file(file_data: bytes, filename: str) -> Tuple[str, dict]:
    log_info(f'Parsing file: {filename}', context='parse_resume_file')

    # Phase 1: Validate
    try:
        is_valid, error_msg, file_type = validate_file(file_data, filename)
        if not is_valid:
            log_warning(f'Validation failed for {filename}: {error_msg}', context='parse_resume_file')
            raise FileValidationError(error_msg)
    except FileValidationError:
        raise
    except Exception as e:
        log_error(e, context='parse_resume_file_validation')
        raise FileValidationError(
            'Could not validate the uploaded file. Please ensure it is a valid PDF or DOCX.'
        ) from e

    # Phase 2: Extract
    try:
        text = extract_text(file_data, file_type)
        log_info(f'Extracted {len(text)} chars from {filename}', context='parse_resume_file')
    except FileParsingError:
        raise
    except Exception as e:
        log_error(e, context='parse_resume_file_extraction')
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    metadata = {
        'filename':        filename,
        'file_type':       file_type,
        'file_size_bytes': len(file_data),
        'text_length':     len(text),
        'success':         True,
    }
    return text, metadata
